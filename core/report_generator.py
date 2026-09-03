import io
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import xlsxwriter

class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 14)
        self.set_text_color(0, 51, 102) # Azul oscuro
        self.cell(0, 10, 'Sistema de Mantenimiento Predictivo IS-402', border=False, align='C')
        self.ln(15)
        
    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Documento Generado Automáticamente - Página {self.page_no()}', align='C')

def generate_pdf(title, summary, df, img_path=None):
    pdf = PDF()
    pdf.add_page()
    
    # Título Principal
    pdf.set_font('helvetica', 'B', 18)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, title, ln=True, align='C')
    pdf.ln(5)
    
    # Resumen Ejecutivo
    pdf.set_font('helvetica', 'B', 12)
    pdf.cell(0, 10, 'Resumen Ejecutivo', ln=True)
    pdf.set_font('helvetica', '', 11)
    pdf.multi_cell(0, 8, summary)
    pdf.ln(10)
    
    # Visualización (Gráfico)
    if img_path:
        pdf.set_font('helvetica', 'B', 12)
        pdf.cell(0, 10, 'Visualización de Datos', ln=True)
        try:
            pdf.image(img_path, w=170)
        except Exception as e:
            pdf.set_font('helvetica', 'I', 10)
            pdf.cell(0, 10, f"[No se pudo incrustar la imagen: {e}]")
        pdf.ln(10)
        
    # Datos Detallados (Tabla)
    pdf.set_font('helvetica', 'B', 12)
    pdf.cell(0, 10, 'Tabla de Datos Detallados', ln=True)
    pdf.set_font('helvetica', '', 8)
    
    if not df.empty:
        # Calcular ancho de columnas dinámico
        num_cols = len(df.columns)
        col_w = 190 / max(num_cols, 1)
        
        # Encabezados
        pdf.set_fill_color(79, 129, 189) # Azul
        pdf.set_text_color(255, 255, 255)
        for col in df.columns:
            pdf.cell(col_w, 8, str(col)[:15], border=1, align='C', fill=True)
        pdf.ln()
        
        # Filas
        pdf.set_text_color(0, 0, 0)
        for index, row in df.head(45).iterrows(): # Limitar a 45 para PDF
            for val in row:
                pdf.cell(col_w, 8, str(val)[:15], border=1, align='C')
            pdf.ln()
            
        if len(df) > 45:
            pdf.set_font('helvetica', 'I', 9)
            pdf.cell(0, 8, f"... y {len(df)-45} registros más omitidos por formato.", border=0)
            
    # Conclusión
    pdf.ln(10)
    pdf.set_font('helvetica', 'B', 12)
    pdf.cell(0, 10, 'Conclusiones', ln=True)
    pdf.set_font('helvetica', '', 11)
    pdf.multi_cell(0, 8, "Este reporte certifica el estado operativo según los filtros solicitados por el usuario. "
                         "Los datos provienen del Data Warehouse en PostgreSQL del sistema minero.")
        
    return bytes(pdf.output())

def generate_word(title, summary, df, img_path=None):
    output = io.BytesIO()
    doc = Document()
    
    # Portada y Título
    doc.add_heading(title, 0)
    
    # Resumen
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph(summary)
    
    # Imagen
    if img_path:
        doc.add_heading('Visualización de Datos', level=1)
        try:
            doc.add_picture(img_path, width=Inches(6.0))
        except:
            doc.add_paragraph("[Error incrustando la imagen. Kaleido o permisos fallaron]")
            
    # Tabla
    doc.add_heading('Tabla de Datos Detallados', level=1)
    if not df.empty:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            
        for index, row in df.head(100).iterrows(): # Límite 100 filas
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
    doc.add_page_break()
    doc.add_paragraph("Generado automáticamente por el Sistema de Mantenimiento Predictivo IS-402.")
    
    doc.save(output)
    return output.getvalue()

def generate_excel(title, summary, df):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet('Reporte Oficial')
    
    # Definición de formatos profesionales
    title_format = workbook.add_format({'bold': True, 'font_size': 16, 'color': '#003366'})
    header_format = workbook.add_format({
        'bold': True, 'bg_color': '#4F81BD', 'font_color': 'white',
        'border': 1, 'align': 'center', 'valign': 'vcenter'
    })
    cell_format = workbook.add_format({'border': 1, 'align': 'left'})
    
    # Escribir Título y Resumen
    worksheet.write(0, 0, title, title_format)
    worksheet.write(2, 0, "Resumen Ejecutivo:", workbook.add_format({'bold': True}))
    worksheet.write(3, 0, summary)
    
    # Escribir DataFrame
    start_row = 6
    if not df.empty:
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(start_row, col_num, value, header_format)
            
        for row_num, row_data in enumerate(df.values):
            for col_num, value in enumerate(row_data):
                worksheet.write(start_row + row_num + 1, col_num, str(value), cell_format)
                
        # Ajuste inteligente de ancho de columnas
        for i, col in enumerate(df.columns):
            max_len = max(df[col].astype(str).map(len).max(), len(str(col))) + 2
            worksheet.set_column(i, i, min(max_len, 40)) # Cap a 40 de ancho
            
    workbook.close()
    return output.getvalue()
